import asyncio
import os
import threading
import tkinter as tk
from asyncio import CancelledError
from tkinter import messagebox, simpledialog
from PIL import Image, ImageTk, ImageDraw, ImageFont
from telethon.errors import SessionPasswordNeededError

from docx import Document
from docx.shared import Pt, RGBColor

from app.utils.constants import AVATAR_SIZE, SESSIONS_DIR, IMAGES_DIR
from app.utils.file_utils import load_meta, save_meta
from app.utils.image_utils import make_rounded_avatar, generate_letter_avatar
from app.telegram_client.client_manager import TelegramClientManager
from app.services.session_service import remove_session


class TelegramLoginApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Telegram Login")
        self.root.geometry("400x500")
        self.root.resizable(False, False)

        # создаём ОДИН event loop
        self.loop = asyncio.new_event_loop()
        threading.Thread(target=self._start_loop, daemon=True).start()

        # передаём loop менеджеру
        self.client_manager = TelegramClientManager(loop=self.loop)
        self.client = None

        self.show_session_selector()

    def _start_loop(self):
        asyncio.set_event_loop(self.loop)
        self.loop.run_forever()

    # -------------------- Helpers --------------------
    def clear_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()

    def disconnect_client(self):
        if self.client_manager:
            self.client_manager.disconnect()

    def get_account_image(self, session_name, display_name):
        meta = load_meta()
        info = meta.get(session_name + ".session", {})
        avatar_path = info.get("avatar")

        if avatar_path and os.path.exists(avatar_path):
            img = Image.open(avatar_path)
            return ImageTk.PhotoImage(make_rounded_avatar(img))
        else:
            img = generate_letter_avatar(display_name[0] if display_name else "?")
            return ImageTk.PhotoImage(img)

    # -------------------- Session Selector --------------------
    def show_session_selector(self):
        self.clear_window()
        self.root.geometry("400x500")
        tk.Label(self.root, text="Choose Telegram Account", font=("Arial", 12, "bold")).pack(pady=10)

        meta = load_meta()
        sessions = [f[:-8] for f in os.listdir(SESSIONS_DIR) if f.endswith(".session") and f != "temp.session"]

        if sessions:
            for s in sessions:
                info = meta.get(s + ".session", {})
                display_name = info.get("display_name", s)
                img = self.get_account_image(s, display_name)

                frame = tk.Frame(self.root)
                frame.pack(pady=5, fill=tk.X, padx=10)

                if img:
                    lbl_img = tk.Label(frame, image=img)
                    lbl_img.image = img
                    lbl_img.pack(side=tk.LEFT, padx=5)

                tk.Button(frame, text=display_name, width=15,
                          command=lambda name=s: self.login_with_existing(name)).pack(side=tk.LEFT)
                tk.Button(frame, text="Remove", fg="red",
                          command=lambda name=s: self.remove_account(name)).pack(side=tk.LEFT, padx=5)
        else:
            tk.Label(self.root, text="No saved accounts found.").pack(pady=10)

        tk.Button(self.root, text="+ Add New Account", command=self.create_api_form).pack(pady=20)

    def remove_account(self, session_name):
        if messagebox.askyesno("Confirm", f"Remove account '{session_name}'?"):
            self.disconnect_client()
            remove_session(session_name)
            messagebox.showinfo("Removed", f"Account '{session_name}' removed.")
            self.show_session_selector()

    # -------------------- Existing Account Login --------------------
    def login_with_existing(self, session_name):
        meta = load_meta()
        info = meta.get(session_name + ".session")

        if not info:
            messagebox.showerror("Error", "API credentials for this account are missing.")
            return

        try:
            self.disconnect_client()
            self.client = self.client_manager.connect(session_name, info["api_id"], info["api_hash"])

            if self.client_manager.is_authorized():
                self.check_and_update_avatar(session_name)
                self.show_success()
            else:
                messagebox.showerror("Error", "Session not authorized. Please log in again.")
                self.create_api_form()
        except Exception as e:
            messagebox.showerror("Error", str(e))

    # -------------------- API Login --------------------
    def create_api_form(self):
        self.clear_window()

        tk.Label(self.root, text="Enter Telegram API Credentials", font=("Arial", 12, "bold")).pack(pady=10)

        tk.Label(self.root, text="API ID:").pack(pady=5)
        tk.Entry(self.root, textvariable=self.api_id).pack(pady=5)

        tk.Label(self.root, text="API Hash:").pack(pady=5)
        tk.Entry(self.root, textvariable=self.api_hash).pack(pady=5)

        # Frame для кнопок
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=15)  # отступ сверху/снизу всего блока кнопок

        tk.Button(button_frame, text="Back to Account List", command=self.show_session_selector).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Next", command=self.start_login).pack(side=tk.LEFT, padx=5)


    def start_login(self):
        try:
            api_id = int(self.api_id.get())
            api_hash = self.api_hash.get().strip()
        except ValueError:
            messagebox.showerror("Error", "API ID must be a number.")
            return

        self.client = self.client_manager.connect("temp", api_id, api_hash)

        if not self.client_manager.is_authorized():
            messagebox.showinfo("Success", "API credentials accepted!")
            self.create_phone_form()
        else:
            self.rename_session_after_login()
            self.show_success()

    # -------------------- Phone Verification --------------------
    def create_phone_form(self):
        self.clear_window()

        tk.Label(self.root, text="Phone number (+country code):").pack(pady=5)
        tk.Entry(self.root, textvariable=self.phone).pack(pady=5)

        # Frame для кнопок
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=15)

        # Кнопка Back слева
        tk.Button(button_frame, text="Back", command=self.create_api_form).pack(side=tk.LEFT, padx=5)
        # Кнопка Send Code справа
        tk.Button(button_frame, text="Send Code", command=self.send_code).pack(side=tk.LEFT, padx=5)

    def send_code(self):
        phone = self.phone.get().strip()
        if not phone:
            messagebox.showerror("Error", "Please enter your phone number.")
            return
        try:
            result = self.client_manager.send_code(phone)
            self.phone_code_hash = result.phone_code_hash
            messagebox.showinfo("Code Sent", "Check your Telegram for the login code.")
            self.create_code_form()
        except Exception as e:
            messagebox.showerror("Send Code Error", str(e))

    def create_code_form(self):
        self.clear_window()

        tk.Label(self.root, text="Enter the code you received in Telegram:").pack(pady=5)
        tk.Entry(self.root, textvariable=self.code).pack(pady=5)

        # Frame для кнопок
        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=15)

        # Кнопка Back слева
        tk.Button(button_frame, text="Back", command=self.create_phone_form).pack(side=tk.LEFT, padx=5)
        # Кнопка Verify Code справа
        tk.Button(button_frame, text="Verify Code", command=self.verify_code).pack(side=tk.LEFT, padx=5)

    def verify_code(self):
        phone = self.phone.get().strip()
        code = self.code.get().strip()
        try:
            self.client_manager.sign_in(phone, code, self.phone_code_hash)
            self.rename_session_after_login()
            self.show_success()
        except SessionPasswordNeededError:
            pw = simpledialog.askstring("2FA Password", "Enter your two-step password:", show="*")
            self.client_manager.sign_in_2fa(pw)
            self.rename_session_after_login()
            self.show_success()
        except Exception as e:
            import traceback
            traceback.print_exc()
            messagebox.showerror("Login Error", str(e))

    # -------------------- Session Rename & Avatar Handling --------------------
    def rename_session_after_login(self):
        me = self.client_manager.get_me()
        name = me.username or me.first_name or str(me.phone)
        safe_name = name.replace(" ", "_").replace("@", "")
        new_path = os.path.join(SESSIONS_DIR, f"{safe_name}.session")

        temp_path = self.temp_session_path + ".session"
        if os.path.exists(temp_path):
            self.disconnect_client()
            os.rename(temp_path, new_path)

        # Save API credentials and avatar
        meta = load_meta()
        meta[safe_name + ".session"] = {
            "api_id": int(self.api_id.get()),
            "api_hash": self.api_hash.get(),
            "display_name": me.first_name or safe_name
        }
        save_meta(meta)

        self.disconnect_client()
        self.client = self.client_manager.connect(safe_name, int(self.api_id.get()), self.api_hash.get())
        self.check_and_update_avatar(safe_name)

    def check_and_update_avatar(self, session_name):
        me = self.client_manager.get_me()
        safe_name = session_name
        photo_path = None

        try:
            if me.photo:
                filename = f"{safe_name}.png"
                photo_path = self.client_manager.download_avatar(me, filename)
            else:
                meta = load_meta()
                info = meta.get(safe_name + ".session")
                if info:
                    old_photo = info.get("avatar")
                    if old_photo and os.path.exists(old_photo):
                        os.remove(old_photo)
        except:
            pass

        meta = load_meta()
        info = meta.get(safe_name + ".session", {})
        info["avatar"] = photo_path
        meta[safe_name + ".session"] = info
        save_meta(meta)

    def send_test_message(self):
        if self.client:
            self.client_manager.send_message("me", "Hello from Tkinter GUI!")
            messagebox.showinfo("Message Sent", "Message sent to Saved Messages.")

    def show_success(self):
        import os
        from PIL import Image, ImageTk, ImageDraw, ImageFont

        # -------------------- Плейсхолдер для отсутствующих аватаров --------------------
        def generate_placeholder_avatar(letter):
            img = Image.new("RGB", (40, 40), color="#cccccc")
            draw = ImageDraw.Draw(img)
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()

            # Новый способ вычислить размеры текста
            bbox = draw.textbbox((0, 0), letter, font=font)
            w = bbox[2] - bbox[0]
            h = bbox[3] - bbox[1]

            # Центрируем текст
            draw.text(((40 - w) / 2, (40 - h) / 2), letter, fill="white", font=font)
            return img

        # -------------------- Основной код --------------------
        # Увеличиваем окно
        self.root.geometry("800x500")

        self.clear_window()
        me = self.client_manager.get_me()

        meta = load_meta()
        session_name = me.username or me.first_name or str(me.phone)
        safe_name = session_name.replace(" ", "_").replace("@", "")
        avatar_path = meta.get(safe_name + ".session", {}).get("avatar")

        # -------------------- Главный контейнер --------------------
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill="both", expand=True)

        # -------------------- Левая панель (сайдбар) --------------------
        sidebar = tk.Frame(main_frame, width=250, bg="#f0f0f0", relief="ridge", bd=2)
        sidebar.pack(side="left", fill="y")

        # Загружаем или генерируем аватар
        if avatar_path and os.path.exists(avatar_path):
            img = Image.open(avatar_path)
            photo = ImageTk.PhotoImage(make_rounded_avatar(img))
        else:
            img = generate_letter_avatar(me.first_name or "?")
            photo = ImageTk.PhotoImage(img)

        lbl_img = tk.Label(sidebar, image=photo, bg="#f0f0f0")
        lbl_img.image = photo
        lbl_img.pack(pady=10)

        tk.Label(sidebar, text=f"{me.first_name}", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(pady=2)
        if me.username:
            tk.Label(sidebar, text=f"@{me.username}", bg="#f0f0f0").pack(pady=2)
        tk.Label(sidebar, text=f"📱 {me.phone}", bg="#f0f0f0").pack(pady=2)

        tk.Button(sidebar, text="📤 Send Test Message", command=self.send_test_message).pack(pady=15, fill="x", padx=10)
        tk.Button(sidebar, text="⬅️ Back to Accounts", command=self.show_session_selector).pack(pady=5, fill="x",
                                                                                                padx=10)

        # -------------------- Правая панель (список диалогов) --------------------
        import asyncio

        # Получаем диалоги (блокирующе, но безопасно)
        future = asyncio.run_coroutine_threadsafe(
            self.client_manager.get_dialogs(None),
            self.loop
        )
        dialogs = future.result()

        dialogs_frame = tk.Frame(main_frame, bg="white")
        dialogs_frame.pack(side="right", fill="both", expand=True)

        tk.Label(dialogs_frame, text="Dialogs", font=("Arial", 13, "bold"), bg="white").pack(anchor="w", padx=10,
                                                                                             pady=5)

        # Canvas + Scrollbar
        canvas = tk.Canvas(dialogs_frame, bg="white", highlightthickness=0)
        scrollbar = tk.Scrollbar(dialogs_frame, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)

        scrollable_frame = tk.Frame(canvas, bg="white")
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        # Обновление области прокрутки при добавлении элементов
        def on_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        scrollable_frame.bind("<Configure>", on_configure)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def _on_mousewheel(event):
            canvas_height = canvas.winfo_height()
            content_height = scrollable_frame.winfo_height()

            if content_height > canvas_height:
                canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        # Привязываем колесо мыши
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        self.selected_dialog_id = None  # ID выбранного диалога
        self.dialog_labels = []  # Список Label для управления цветом
        DIALOG_BG = "white"
        DIALOG_BG_SELECTED = "#cce5ff"  # цвет выбранного диалога

        def select_dialog(dialog_id, label):
            # Сбрасываем фон всех
            for lbl in self.dialog_labels:
                lbl.config(bg=DIALOG_BG)
            # Выделяем выбранный
            label.config(bg=DIALOG_BG_SELECTED)
            self.selected_dialog_id = dialog_id
            print("Выбран диалог ID:", dialog_id)

        # Выводим диалоги
        for d in dialogs:
            print(d)

            # Генерируем плейсхолдер-аватар с первой буквой названия
            first_letter = (d.name[0].upper() if d.name else "?")
            avatar_img = generate_placeholder_avatar(first_letter)
            avatar_photo = ImageTk.PhotoImage(avatar_img)

            # Контейнер для одной строки диалога
            dialog_frame = tk.Frame(scrollable_frame, bg=DIALOG_BG, padx=5, pady=3)
            dialog_frame.pack(fill="x", padx=10, pady=2)

            # Аватар
            avatar_label = tk.Label(dialog_frame, image=avatar_photo, bg=DIALOG_BG)
            avatar_label.image = avatar_photo  # нужно, чтобы не сборщик мусора не удалил
            avatar_label.pack(side="left", padx=(0, 8))

            # Название диалога
            lbl = tk.Label(
                dialog_frame,
                text=f"{d.name}",
                bg=DIALOG_BG,
                anchor="w",
                font=("Arial", 11),
                cursor="hand2"
            )
            lbl.pack(side="left", fill="x", expand=True)

            # Клик по строке диалога
            def on_click(event, dialog_id=d.id, l=dialog_frame):
                select_dialog(dialog_id, l)

            dialog_frame.bind("<Button-1>", on_click)
            lbl.bind("<Button-1>", on_click)
            avatar_label.bind("<Button-1>", on_click)

            self.dialog_labels.append(dialog_frame)

            # -------------------- Панель для экспорта (изначально скрыта) --------------------
            # -------------------- Панель для экспорта (одна, общая, изначально скрыта) --------------------
            self.export_controls = tk.Frame(dialogs_frame, bg="#eef5ff", relief="ridge", bd=2)
            self.export_controls.pack(fill="x", padx=10, pady=(0, 10))
            self.export_controls.pack_forget()  # скрываем до выбора диалога

            export_label = tk.Label(self.export_controls, text="📜", bg="#eef5ff")
            export_label.pack(side="left", padx=5)

            count_entry = tk.Entry(self.export_controls, width=6)
            count_entry.insert(0, "50")
            count_entry.pack(side="left", padx=5)



            export_word_btn = tk.Button(
                self.export_controls,
                text="📄 Export to Word",
                command=lambda: export_chat_to_docx(self.selected_dialog, asyncio.run_coroutine_threadsafe(
                    self.client_manager.client.get_messages(self.selected_dialog, limit=int(count_entry.get())),
                    self.loop
                ).result())
            )
            export_word_btn.pack(side="left", padx=10)

            # -------------------- Функции --------------------

            def select_dialog(dialog_id, frame):
                # Сброс выделений
                for lbl in self.dialog_labels:
                    lbl.config(bg=DIALOG_BG)
                frame.config(bg=DIALOG_BG_SELECTED)

                self.selected_dialog_id = dialog_id
                self.selected_dialog = next((d for d in dialogs if d.id == dialog_id), None)
                print("Выбран диалог ID:", dialog_id)

                # Показываем панель экспорта
                self.export_controls.pack(fill="x", padx=10, pady=(0, 10))

            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            import asyncio

            def export_chat_to_docx(dialog, messages):
                doc = Document()
                doc.add_heading(f"💬 Chat with {dialog.name}", level=1)

                # Получаем текущего пользователя (чтобы понимать, какие сообщения — свои)
                me = asyncio.run_coroutine_threadsafe(
                    self.client_manager.client.get_me(), self.loop
                ).result()

                for msg in reversed(messages):
                    sender = getattr(msg.sender, "first_name", "Unknown")
                    text = msg.message or ""
                    time_str = msg.date.strftime("%Y-%m-%d %H:%M")
                    is_me = (msg.sender_id == me.id)

                    # Основной контейнер параграфа
                    p = doc.add_paragraph()

                    # ✅ Имя отправителя (жирным)
                    sender_run = p.add_run(sender)
                    sender_run.bold = True
                    sender_run.font.size = Pt(11)
                    sender_run.font.color.rgb = RGBColor(0, 102, 204) if is_me else RGBColor(0, 0, 0)

                    # ✅ Время под именем (курсив, тонкое начертание)
                    p.add_run("\n")  # переход на новую строку
                    time_run = p.add_run(time_str)
                    time_run.italic = True
                    time_run.font.size = Pt(8)
                    time_run.font.color.rgb = RGBColor(128, 128, 128)

                    # ✅ Текст сообщения
                    p.add_run("\n")
                    text_run = p.add_run(text)
                    text_run.font.size = Pt(11)
                    text_run.font.color.rgb = RGBColor(0, 0, 0)

                    # ✅ Выровнять по левой/правой стороне
                    if is_me:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # свои сообщения — справа
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # чужие — слева

                    # Добавляем отступ между сообщениями
                    doc.add_paragraph("")

                # Сохраняем DOCX
                os.makedirs("exports/docx", exist_ok=True)
                file_path = f"exports/docx/chat_{dialog.id}.docx"
                doc.save(file_path)
                print(f"✅ Экспортировано в Word: {file_path}")





